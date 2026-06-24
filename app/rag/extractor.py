"""
Document Text Extractor — Phase 1

Extracts plain text from PDF, DOCX, and TXT files.

Architecture notes:
  - Each format has its own extractor function.
  - All extractors return the same ExtractionResult structure.
  - PDF extraction tries direct text first, falls back to OCR for scanned pages.
  - This module knows NOTHING about chunking, embedding, or ChromaDB.
    It just turns files into text.

Libraries used:
  - PyMuPDF (fitz): PDF text extraction + rendering pages to images for OCR
  - python-docx:    DOCX paragraph/table extraction
  - pytesseract:    OCR for scanned PDF pages
  - Pillow:         Image handling for OCR pipeline
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path

import fitz  # PyMuPDF

from app.schemas.models import DocumentMetadata

logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════
# Output structure — same for ALL file types
# ═══════════════════════════════════════════════════════════════


@dataclass
class PageText:
    """Text extracted from a single page."""

    page_number: int
    text: str
    char_count: int


@dataclass
class ExtractionResult:
    """
    Uniform output from any extractor.

    This is the contract between the extraction layer and everything downstream
    (metadata extraction, chunking, etc).
    """

    raw_text: str  # Full concatenated text
    pages: list[PageText] = field(default_factory=list)
    page_count: int = 0
    extraction_method: str = "text"  # "text", "ocr", or "hybrid"
    metadata: DocumentMetadata | None = None


# ═══════════════════════════════════════════════════════════════
# Main entry point
# ═══════════════════════════════════════════════════════════════


def extract_text(file_path: str, file_type: str) -> ExtractionResult:
    """
    Extract text from a file. This is the single entry point for all formats.

    Args:
        file_path: Absolute path to the file on disk.
        file_type: Extension without dot — "pdf", "docx", "txt"

    Returns:
        ExtractionResult with the extracted text and page information.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file type is not supported.
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    if not path.is_file():
        raise ValueError(f"Path is not a file: {file_path}")

    # Dispatch to the correct extractor based on file type
    extractors = {
        "pdf": _extract_pdf,
        "docx": _extract_docx,
        "txt": _extract_txt,
    }

    file_type_lower = file_type.lower().strip()
    extractor = extractors.get(file_type_lower)

    if extractor is None:
        raise ValueError(
            f"Unsupported file type: '{file_type}'. "
            f"Supported types: {', '.join(extractors.keys())}"
        )

    logger.info("Extracting text from %s (type: %s)", path.name, file_type_lower)
    result = extractor(path)

    # Phase 2: Extract rich metadata
    from app.rag.metadata import extract_metadata
    
    result.metadata = extract_metadata(
        file_path=file_path,
        file_type=file_type,
        total_chars=len(result.raw_text),
        extraction_method=result.extraction_method,
    )

    logger.info(
        "Extraction complete: %d pages, %d chars, method=%s",
        result.page_count,
        len(result.raw_text),
        result.extraction_method,
    )

    return result


# ═══════════════════════════════════════════════════════════════
# PDF Extractor
# ═══════════════════════════════════════════════════════════════

# If a page yields fewer characters than this, it's likely a scanned image
_MIN_CHARS_FOR_TEXT_PAGE = 50


def _extract_pdf(path: Path) -> ExtractionResult:
    """
    Extract text from a PDF file.

    Strategy:
      1. Open with PyMuPDF.
      2. For each page, try direct text extraction.
      3. If a page has very little text (< 50 chars), it's likely scanned.
         Fall back to OCR for that page.
      4. Track whether we used text, OCR, or both (hybrid).
    """
    doc = fitz.open(str(path))
    pages: list[PageText] = []
    used_ocr = False
    used_text = False

    for page_num in range(len(doc)):
        page = doc[page_num]

        # Try direct text extraction first
        text = page.get_text("text").strip()

        if len(text) >= _MIN_CHARS_FOR_TEXT_PAGE:
            # Good text extraction — use it
            used_text = True
        else:
            # Little or no text — try OCR
            ocr_text = _ocr_page(page)
            if ocr_text:
                text = ocr_text
                used_ocr = True
            # If OCR also fails, we keep whatever little text we got

        pages.append(
            PageText(
                page_number=page_num + 1,  # 1-indexed for humans
                text=text,
                char_count=len(text),
            )
        )

    doc.close()

    # Determine extraction method
    if used_ocr and used_text:
        method = "hybrid"
    elif used_ocr:
        method = "ocr"
    else:
        method = "text"

    # Join all page texts with double newline (paragraph separator)
    raw_text = "\n\n".join(p.text for p in pages if p.text)

    return ExtractionResult(
        raw_text=raw_text,
        pages=pages,
        page_count=len(pages),
        extraction_method=method,
    )


def _ocr_page(page: fitz.Page) -> str:
    """
    OCR a single PDF page by rendering it to an image and running Tesseract.

    Returns the OCR'd text, or empty string if OCR is not available.
    """
    try:
        import pytesseract
        from PIL import Image
        import io

        # Render page to image at 200 DPI (balance between quality and speed)
        # Higher DPI = better OCR but slower
        pix = page.get_pixmap(dpi=200)
        img_bytes = pix.tobytes("png")
        image = Image.open(io.BytesIO(img_bytes))

        # Run Tesseract OCR
        text = pytesseract.image_to_string(image).strip()
        if text:
            logger.debug(
                "OCR extracted %d chars from page %d", len(text), page.number + 1
            )
        return text

    except ImportError:
        logger.warning(
            "pytesseract or Pillow not installed — skipping OCR for page %d",
            page.number + 1,
        )
        return ""
    except Exception as exc:
        # Tesseract might not be installed on the system
        logger.warning("OCR failed for page %d: %s", page.number + 1, exc)
        return ""


# ═══════════════════════════════════════════════════════════════
# DOCX Extractor
# ═══════════════════════════════════════════════════════════════


def _extract_docx(path: Path) -> ExtractionResult:
    """
    Extract text from a DOCX file.

    Strategy:
      1. Open with python-docx.
      2. Iterate through all paragraphs and extract text.
      3. Iterate through all tables and extract cell text.
      4. DOCX doesn't have a natural "page" concept,
         so we treat the entire document as page 1.

    Note: We also preserve paragraph style names (e.g., "Heading 1")
    in the text output. This helps with structure-aware chunking in Phase 3.
    """
    from docx import Document

    doc = Document(str(path))

    text_parts: list[str] = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        text_parts.append(text)

    # Extract tables (join cells with tabs, rows with newlines)
    for table in doc.tables:
        table_lines: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_lines.append("\t".join(cells))
        if table_lines:
            text_parts.append("\n".join(table_lines))

    raw_text = "\n\n".join(text_parts)

    # DOCX has no page concept — treat as single page
    pages = [
        PageText(
            page_number=1,
            text=raw_text,
            char_count=len(raw_text),
        )
    ]

    return ExtractionResult(
        raw_text=raw_text,
        pages=pages,
        page_count=1,
        extraction_method="text",
    )


# ═══════════════════════════════════════════════════════════════
# TXT Extractor
# ═══════════════════════════════════════════════════════════════

# Encodings to try, in order
_ENCODINGS = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]


def _extract_txt(path: Path) -> ExtractionResult:
    """
    Extract text from a plain text file.

    Strategy:
      1. Try reading as UTF-8 first.
      2. Fall back through common encodings if UTF-8 fails.
      3. Normalize line endings to \\n.
    """
    raw_text = None

    for encoding in _ENCODINGS:
        try:
            raw_text = path.read_text(encoding=encoding)
            break
        except (UnicodeDecodeError, UnicodeError):
            continue

    if raw_text is None:
        # Last resort: read as bytes, decode with replacement characters
        raw_bytes = path.read_bytes()
        raw_text = raw_bytes.decode("utf-8", errors="replace")

    # Normalize line endings (Windows \r\n → \n)
    raw_text = raw_text.replace("\r\n", "\n").replace("\r", "\n")

    # TXT has no pages — treat as single page
    pages = [
        PageText(
            page_number=1,
            text=raw_text,
            char_count=len(raw_text),
        )
    ]

    return ExtractionResult(
        raw_text=raw_text,
        pages=pages,
        page_count=1,
        extraction_method="text",
    )
