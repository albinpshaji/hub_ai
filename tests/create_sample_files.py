"""
Generate sample PDF and DOCX test files for the extraction pipeline.

Run with:
    cd hub_ai
    .\venv\Scripts\python tests\create_sample_files.py
"""

from pathlib import Path


SAMPLE_DIR = Path(__file__).parent / "sample_files"
SAMPLE_DIR.mkdir(exist_ok=True)


def create_sample_pdf():
    """Create a multi-page PDF with text content using PyMuPDF."""
    import fitz

    doc = fitz.open()  # new empty PDF

    pages_content = [
        {
            "title": "Chapter 1: Introduction to Machine Learning",
            "body": (
                "Machine learning is a branch of artificial intelligence that focuses on "
                "building systems that learn from data. Unlike traditional programming where "
                "rules are explicitly coded, machine learning algorithms build models based "
                "on sample data, known as training data, to make predictions or decisions "
                "without being explicitly programmed.\n\n"
                "The field of machine learning has grown rapidly in recent years, driven by "
                "advances in computing power, the availability of large datasets, and "
                "improvements in algorithms. Today, machine learning powers many applications "
                "we use daily, from email spam filters to recommendation systems on streaming "
                "platforms."
            ),
        },
        {
            "title": "Chapter 2: Types of Machine Learning",
            "body": (
                "There are three main types of machine learning:\n\n"
                "1. Supervised Learning: The algorithm learns from labeled training data, "
                "finding patterns that map inputs to known outputs. Common examples include "
                "classification (spam detection) and regression (price prediction).\n\n"
                "2. Unsupervised Learning: The algorithm finds hidden patterns in data without "
                "labeled outputs. Clustering and dimensionality reduction are key techniques.\n\n"
                "3. Reinforcement Learning: The agent learns by interacting with an environment, "
                "receiving rewards or penalties. This approach is used in game playing, robotics, "
                "and autonomous vehicles."
            ),
        },
        {
            "title": "Chapter 3: Neural Networks and Deep Learning",
            "body": (
                "Neural networks are computing systems inspired by biological neural networks "
                "in the human brain. A neural network consists of layers of interconnected "
                "nodes (neurons) that process information.\n\n"
                "Deep learning uses neural networks with many hidden layers (hence 'deep'). "
                "The key architectures include:\n\n"
                "- Convolutional Neural Networks (CNNs): Specialized for image processing\n"
                "- Recurrent Neural Networks (RNNs): Designed for sequential data\n"
                "- Transformers: The architecture behind modern language models like GPT and BERT\n\n"
                "The transformer architecture, introduced in the paper 'Attention Is All You Need' "
                "by Vaswani et al. (2017), revolutionized natural language processing and has "
                "become the foundation for large language models."
            ),
        },
    ]

    for page_data in pages_content:
        page = doc.new_page(width=595, height=842)  # A4 size

        # Title
        title_rect = fitz.Rect(50, 50, 545, 100)
        page.insert_textbox(
            title_rect,
            page_data["title"],
            fontsize=16,
            fontname="helv",
            color=(0, 0, 0.6),
        )

        # Body
        body_rect = fitz.Rect(50, 120, 545, 792)
        page.insert_textbox(
            body_rect,
            page_data["body"],
            fontsize=11,
            fontname="helv",
        )

    output_path = SAMPLE_DIR / "sample.pdf"
    doc.save(str(output_path))
    doc.close()
    print(f"Created: {output_path} ({output_path.stat().st_size} bytes)")


def create_sample_docx():
    """Create a DOCX with headings, paragraphs, and a table."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # Title
    doc.add_heading("Introduction to Natural Language Processing", level=0)

    # Section 1
    doc.add_heading("What is NLP?", level=1)
    doc.add_paragraph(
        "Natural Language Processing (NLP) is a subfield of linguistics, computer science, "
        "and artificial intelligence concerned with the interactions between computers and "
        "human language. The goal is to enable computers to understand, interpret, and "
        "generate human language in a valuable way."
    )

    # Section 2
    doc.add_heading("Key NLP Tasks", level=1)
    doc.add_paragraph(
        "NLP encompasses a wide range of tasks, from simple text classification to complex "
        "language generation. Some of the most important tasks include:"
    )

    doc.add_heading("Text Classification", level=2)
    doc.add_paragraph(
        "Text classification assigns predefined categories to text documents. "
        "Applications include sentiment analysis, spam detection, and topic categorization."
    )

    doc.add_heading("Named Entity Recognition", level=2)
    doc.add_paragraph(
        "Named Entity Recognition (NER) identifies and classifies named entities in text "
        "into predefined categories such as person names, organizations, locations, and dates."
    )

    doc.add_heading("Machine Translation", level=2)
    doc.add_paragraph(
        "Machine translation automatically translates text from one language to another. "
        "Modern systems use neural machine translation with transformer architectures."
    )

    # Table
    doc.add_heading("Comparison of NLP Approaches", level=1)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"

    headers = ["Approach", "Strengths", "Weaknesses"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    data = [
        ["Rule-based", "Precise, interpretable", "Brittle, hard to scale"],
        ["Statistical", "Data-driven, adaptable", "Needs large datasets"],
        ["Deep Learning", "State-of-the-art accuracy", "Computationally expensive"],
    ]
    for row_idx, row_data in enumerate(data, 1):
        for col_idx, cell_data in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = cell_data

    output_path = SAMPLE_DIR / "sample.docx"
    doc.save(str(output_path))
    print(f"Created: {output_path} ({output_path.stat().st_size} bytes)")


if __name__ == "__main__":
    create_sample_pdf()
    create_sample_docx()
    print("\nAll sample files created successfully!")
